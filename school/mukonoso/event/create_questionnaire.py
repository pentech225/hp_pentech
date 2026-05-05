"""
紙のアンケート Word文書を生成するスクリプト
出力: questionnaire_20260503.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ---- ページ設定（A4、余白狭め） ----
section = doc.sections[0]
section.page_width  = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin    = Cm(1.5)
section.bottom_margin = Cm(1.5)
section.left_margin   = Cm(2.0)
section.right_margin  = Cm(2.0)


# ---- ヘルパー関数 ----

def set_spacing(para, before=0, after=0, line_rule=None, line=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line is not None:
        from docx.shared import Pt as pt
        pf.line_spacing = Pt(line)


def heading(text, size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=4):
    p = doc.add_paragraph()
    p.alignment = align
    set_spacing(p, before=before, after=after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    return p


def subtext(text, size=10, italic=False, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=2):
    p = doc.add_paragraph()
    p.alignment = align
    set_spacing(p, before=before, after=after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.italic = italic
    return p


def question(text, before=10, after=3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p, before=before, after=after)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = True
    return p


def check_option(text, before=1, after=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p, before=before, after=before)
    run = p.add_run(f'　□　{text}')
    run.font.size = Pt(11)
    return p


def writing_line(before=4, after=8):
    """下線付きの記入欄を1行追加"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p, before=before, after=after)
    run = p.add_run('　' * 60)
    run.font.size = Pt(11)
    run.font.underline = True
    return p


def divider(before=6, after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=before, after=after)
    run = p.add_run('─' * 45)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    return p


# ======== 文書本体 ========

# タイトル
heading('プログラミング体験会　アンケート', size=18, before=0, after=2)
subtext('2026年5月3日（日）　武庫西生涯学習プラザ', size=11, before=0, after=2)
subtext('主催：iTeen 子どものためのICT教育普及促進委員会', size=10, italic=True, before=0, after=4)

divider(before=4, after=4)

subtext(
    '本日はご参加いただきまして、誠にありがとうございます。'
    'よりよいイベントのため、ぜひご意見をお聞かせください。',
    size=10, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=6
)

divider(before=2, after=2)

# Q1
question('１．お子さんのお名前をご記入ください。')
writing_line()

# Q2
question('２．今日の体験はいかがでしたか？　当てはまる数字に○をつけてください。')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_spacing(p, before=2, after=2)
run = p.add_run('　　１（つまらなかった）　　　　２　　　　　３（ふつう）　　　　　４　　　　　５（とても楽しかった）')
run.font.size = Pt(10)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_spacing(p2, before=2, after=6)
# 数字を大きく目立たせる
for num in ['　　　　１', '　　　　　　　　２', '　　　　　　　　３', '　　　　　　　　４', '　　　　　　　　５']:
    r = p2.add_run(num)
    r.font.size = Pt(16)
    r.font.bold = True

# Q3
question('３．お子さんの様子を見て、プログラミングを継続して学ばせたいと思いましたか？')
check_option('ぜひそう思った')
check_option('少し思った')
check_option('今は考えていない')

# Q4（追加質問）
question('４．今後も同様のイベントに参加したいですか？また、ご案内を受け取りますか？')
check_option('ぜひ参加したい・案内を受け取りたい')
check_option('興味があれば参加したい')
check_option('今のところ予定はない')

# Q5
question('５．公式LINEにご登録いただけましたか？')
check_option('はい（登録済み）')
check_option('いいえ')
check_option('これからする')

# Q6
question('６．お子さんの写真・動画のSNS等への掲載についてお教えください。')
check_option('OK（顔出し可）')
check_option('OK（後ろ姿のみ）')
check_option('NG（掲載不可）')

# Q7
question('７．ご自由にお書きください（ご感想・ご要望など）')
writing_line()
writing_line(before=2, after=8)
writing_line(before=2, after=8)

divider(before=8, after=4)

# フッター
subtext(
    'ご協力ありがとうございました。アンケートはスタッフにお渡しください。',
    size=9, italic=True, before=0, after=0
)

# ======== 保存 ========
output_path = r'C:\Users\PC\.company\10_hp\hp_pentech\school\mukonoso\event\questionnaire_20260503.docx'
doc.save(output_path)
print(f'✅ 作成完了: {output_path}')
